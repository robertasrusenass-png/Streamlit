import streamlit as st
import requests
import time
from docx import Document
from io import BytesIO


# ==================================================
# PUSLAPIS
# ==================================================

st.title("Audio-to-text")

st.write("Įkelkite garso failą transkribavimui.")


# ==================================================
# KALBOS PASIRINKIMAS
# ==================================================

doc = Document()
doc.add_paragraph(st.session_state["tekstas"])

buffer = BytesIO()
doc.save(buffer)
buffer.seek(0)

kalbos = {
    "Automatiškai (LT / RU)": "auto",
    "Lietuvių": "lt",
    "Rusų": "ru"
}

pasirinkta_kalba = st.selectbox(
    "Pasirinkite garso kalbą",
    list(kalbos.keys())
)


# ==================================================
# DEEPGRAM TRANSKRIBAVIMAS
# ==================================================

def garso_i_teksta(audio_file, kalba):

    url = "https://api.deepgram.com/v1/listen"

    if kalba == "auto":

        params = [
            ("model", "nova-3-general"),
            ("detect_language", "lt"),
            ("detect_language", "ru"),
            ("smart_format", "true")
        ]

    else:

        params = {
            "model": "nova-3-general",
            "language": kalba,
            "smart_format": "true"
        }

    headers = {
        "Authorization": f"Token {st.secrets['DEEPGRAM_API_KEY']}",
        "Content-Type": audio_file.type
    }

    pradzia = time.perf_counter()

    response = requests.post(
        url,
        params=params,
        headers=headers,
        data=audio_file.getvalue(),
        timeout=300
    )

    apdorojimo_laikas = time.perf_counter() - pradzia

    if response.status_code != 200:

        raise Exception(
            f"Deepgram klaida "
            f"{response.status_code}: "
            f"{response.text}"
        )

    rezultatas = response.json()

    channel = rezultatas["results"]["channels"][0]

    alternative = channel["alternatives"][0]

    tekstas = alternative.get(
        "transcript",
        ""
    )

    transkripcijos_patikimumas = alternative.get(
        "confidence",
        None
    )

    # Automatinio kalbos aptikimo rezultatai
    if kalba == "auto":

        aptikta_kalba = channel.get(
            "detected_language",
            "nežinoma"
        )

        kalbos_patikimumas = channel.get(
            "language_confidence",
            None
        )

    else:

        aptikta_kalba = kalba
        kalbos_patikimumas = None

    trukme = rezultatas.get(
        "metadata",
        {}
    ).get(
        "duration",
        0
    )

    return (
        tekstas,
        aptikta_kalba,
        kalbos_patikimumas,
        transkripcijos_patikimumas,
        trukme,
        apdorojimo_laikas
    )


# ==================================================
# FAILO ĮKĖLIMAS
# ==================================================

ikeltas_garso_failas = st.file_uploader(
    "Įkelkite garso failą",
    type=[
        "mp3",
        "wav",
        "m4a",
        "ogg",
        "flac"
    ]
)


if ikeltas_garso_failas is not None:

    st.audio(ikeltas_garso_failas)

    if st.button(
        "Transkribuoti",
        type="primary"
    ):

        with st.spinner(
            "Garso failas transkribuojamas..."
        ):

            try:

                kalbos_kodas = kalbos[
                    pasirinkta_kalba
                ]

                (
                    tekstas,
                    kalba,
                    kalbos_patikimumas,
                    transkripcijos_patikimumas,
                    trukme,
                    apdorojimo_laikas
                ) = garso_i_teksta(
                    ikeltas_garso_failas,
                    kalbos_kodas
                )


                # ==========================================
                # INFORMACIJA
                # ==========================================

                if kalbos_kodas == "auto":

                    st.write(
                        f"Aptikta kalba: "
                        f"**{kalba}**"
                    )

                    if kalbos_patikimumas is not None:

                        st.write(
                            "Kalbos aptikimo patikimumas: "
                            f"**{kalbos_patikimumas:.2%}**"
                        )

                else:

                    kalbos_pavadinimai = {
                        "lt": "Lietuvių",
                        "ru": "Rusų"
                    }

                    st.write(
                        "Pasirinkta kalba: "
                        f"**{kalbos_pavadinimai.get(kalbos_kodas, kalbos_kodas)}**"
                    )


                if transkripcijos_patikimumas is not None:

                    st.write(
                        "Transkripcijos patikimumas: "
                        f"**{transkripcijos_patikimumas:.2%}**"
                    )


                st.write(
                    f"Garso trukmė: "
                    f"**{trukme:.1f} s**"
                )

                st.write(
                    f"Transkribavimo laikas: "
                    f"**{apdorojimo_laikas:.1f} s**"
                )


                # ==========================================
                # REZULTATAS
                # ==========================================

                if tekstas.strip():

                    st.session_state["tekstas"] = tekstas

                    st.success(
                        "Transkribavimas baigtas!"
                    )

                else:

                    st.warning(
                        "Deepgram apdorojo failą, "
                        "bet neatpažino teksto."
                    )


            except Exception as e:

                st.error(
                    f"Įvyko klaida: {e}"
                )


# ==================================================
# TRANSKRIPTO RODYMAS
# ==================================================

if "tekstas" in st.session_state:

    doc = Document()
    doc.add_paragraph(st.session_state["tekstas"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        "Atsisiųsti teksto failą",
        data=st.session_state["tekstas"],
        file_name="transkriptas.txt",
        mime="text/plain",
        key="download_txt"
    )

    st.download_button(
        "Atsisiųsti Word dokumentą",
        data=buffer.getvalue(),
        file_name="transkriptas.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_docx"
    )
